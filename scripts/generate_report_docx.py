from __future__ import annotations

from datetime import date
from pathlib import Path
import platform

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from submission_metrics import (
    AUTHOR_NAME,
    COURSE_NAME,
    DEPARTMENT,
    GITHUB_USERNAME,
    INSTRUCTOR_NAME,
    KEYWORDS,
    PROJECT_TITLE,
    SEMESTER,
    group_rows_by_size,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_FILE = REPO_ROOT / "report" / "final_report_formal.docx"
ASSET_DIR = REPO_ROOT / "report" / "assets"
WORKFLOW_IMAGE = ASSET_DIR / "workflow.png"
PLOT_IMAGE = REPO_ROOT / "experiments" / "results" / "plots.png"
DEMO_STRIP_IMAGE = ASSET_DIR / "demo_strip.png"

BODY_FONT = "TH Sarabun New"
CODE_FONT = "Courier New"


def set_run_font(run, font_name=BODY_FONT, size=12, bold=False, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_document_language(doc: Document):
    styles = doc.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(13)


def set_section_columns(section, columns: int = 2):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_element = cols[0] if cols else OxmlElement("w:cols")
    cols_element.set(qn("w:num"), str(columns))
    cols_element.set(qn("w:space"), "720")
    if not cols:
        sect_pr.append(cols_element)


def set_section_margins(section, margin_inches: float = 0.65):
    section.top_margin = Inches(margin_inches)
    section.bottom_margin = Inches(margin_inches)
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)


def add_paragraph(doc: Document, text: str, font_size=12, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, italic=italic)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 14, 3: 13}[level], bold=True)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)


def format_table(table, header_size=11, body_size=10.5):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=header_size if r_idx == 0 else body_size, bold=(r_idx == 0))


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line)
        set_run_font(run, font_name=CODE_FONT, size=9)


def add_full_width_figure(doc: Document, image_path: Path, caption: str, width_inches=6.3):
    if not image_path.exists():
        return
    pre_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_section_columns(pre_section, columns=1)
    set_section_margins(pre_section, margin_inches=0.6)
    doc.add_picture(str(image_path), width=Inches(width_inches))
    add_caption(doc, caption)
    post_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_section_columns(post_section, columns=2)
    set_section_margins(post_section, margin_inches=0.6)


def compute_metrics():
    grouped = group_rows_by_size()
    speeds = {}
    for size, rows in grouped.items():
        baseline = rows["baseline"]
        proposed = rows["proposed"]
        speed = ((baseline.ms_per_frame - proposed.ms_per_frame) / baseline.ms_per_frame) * 100.0
        jitter_reduction = ((baseline.side_switches - proposed.side_switches) / baseline.side_switches) * 100.0
        memory_reduction = ((baseline.peak_kb - proposed.peak_kb) / baseline.peak_kb) * 100.0
        speeds[size] = {
            "baseline": baseline,
            "proposed": proposed,
            "speedup": speed,
            "jitter_reduction": jitter_reduction,
            "memory_reduction": memory_reduction,
        }
    return grouped, speeds


def add_cover_page(doc: Document, speeds):
    largest = speeds[max(speeds)]
    add_paragraph(doc, PROJECT_TITLE, font_size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "รายงานโครงงานรายวิชา Computer Animation", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, DEPARTMENT, font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"รายวิชา: {COURSE_NAME}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"ผู้จัดทำ: {AUTHOR_NAME}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"GitHub Username: {GITHUB_USERNAME}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"อาจารย์ผู้สอน: {INSTRUCTOR_NAME}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"ภาคการศึกษา: {SEMESTER}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"วันที่จัดทำ: {date.today().strftime('%d %B %Y')}", font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "คำสำคัญ: " + ", ".join(KEYWORDS), font_size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "", font_size=8)
    add_paragraph(
        doc,
        (
            "เอกสารฉบับนี้นำเสนอการออกแบบและประเมินผลระบบติดตามการเคลื่อนไหวของท่าออกกำลังกายแบบเรียลไทม์ "
            "โดยตีความปัญหาในกรอบของ Computer Animation ผ่านโครงกระดูกเชิงท่าทาง (pose skeleton), key poses, "
            "progress interpolation และ temporal state transitions"
        ),
        font_size=13,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        (
            f"สรุปผลเด่นที่ขนาดข้อมูล 100,000 เฟรม: วิธี proposed ลด peak memory จาก "
            f"{largest['baseline'].peak_kb:.3f} KB เป็น {largest['proposed'].peak_kb:.3f} KB "
            f"และลด side-switch jitter จาก {largest['baseline'].side_switches} เป็น "
            f"{largest['proposed'].side_switches} ครั้ง แม้มี runtime overhead เล็กน้อยที่สเกลสูงสุด"
        ),
        font_size=12.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def add_report_body(doc: Document):
    grouped, speeds = compute_metrics()
    m1 = speeds[1_000]
    m2 = speeds[10_000]
    m3 = speeds[100_000]

    add_heading(doc, "บทคัดย่อ (Abstract)", 1)
    add_paragraph(
        doc,
        (
            "โครงงานนี้ศึกษาปัญหาการติดตามการเคลื่อนไหวของท่าออกกำลังกายแบบเรียลไทม์โดยใช้ภาพจากกล้องเว็บแคม "
            "และตีความปัญหาในมุมมองของ Computer Animation แทนการมองเป็นเพียงงานจำแนกภาพจากคอมพิวเตอร์วิทัศน์เท่านั้น "
            "วัตถุประสงค์หลักคือแปลงลำดับของ 2D pose landmarks ที่มีสัญญาณรบกวนให้เป็นสถานะการเคลื่อนไหวที่เสถียร "
            "เช่น count, stage, angle และ progress เพื่อใช้ทั้งในการนับ repetition และแสดงผลเชิงภาพบนระบบแบบ interactive "
            "วิธี baseline ใช้การคำนวณ elbow angle และ threshold แบบตรงไปตรงมาโดยเลือกแขนที่มองเห็นชัดกว่าในแต่ละเฟรม "
            "ส่วนวิธี proposed เพิ่ม confidence-aware side locking และ exponential temporal filtering เพื่อลดความแกว่งของการติดตาม "
            "การทดลองทำบนข้อมูลสังเคราะห์ 3 ขนาด ได้แก่ 1,000, 10,000 และ 100,000 เฟรม พร้อม pipeline สำหรับข้อมูลจริงจาก webcam "
            f"ผลทดลองพบว่าวิธี proposed ยังรักษา repetition-count error เท่ากับศูนย์ในทุกสเกล และเร็วกว่า baseline ที่สเกล 1k และ 10k "
            f"คิดเป็น {m1['speedup']:.2f}% และ {m2['speedup']:.2f}% ตามลำดับ ขณะที่ที่ 100k เฟรมมี runtime overhead เล็กน้อย "
            f"ประมาณ {abs(m3['speedup']):.2f}% อย่างไรก็ตาม วิธี proposed ลด side-switch jitter จาก "
            f"{m3['baseline'].side_switches} เหลือ {m3['proposed'].side_switches} ครั้ง ({m3['jitter_reduction']:.2f}%) "
            f"และลด peak memory จาก {m3['baseline'].peak_kb:.3f} KB เหลือ {m3['proposed'].peak_kb:.3f} KB "
            f"({m3['memory_reduction']:.2f}%) ที่สเกลสูงสุด แสดงให้เห็น trade-off ที่เหมาะสมระหว่างความเร็ว ความเสถียร และความสามารถในการอธิบายผล"
        ),
        font_size=12,
    )

    add_heading(doc, "1. นิยามปัญหา (Problem Definition)", 1)
    add_heading(doc, "1.1 ที่มาและความสำคัญ", 2)
    add_paragraph(
        doc,
        (
            "การติดตามท่าออกกำลังกายแบบเรียลไทม์เป็นโจทย์ที่เหมาะสมอย่างยิ่งสำหรับรายวิชา Computer Animation "
            "เนื่องจากระบบต้องตีความร่างกายมนุษย์เป็นโครงสร้างเชิงข้อต่อ (articulated structure) และต้องตัดสินการเปลี่ยนผ่านระหว่าง key poses "
            "อย่างต่อเนื่องในเชิงเวลา มิใช่เพียงตรวจว่าภาพหนึ่งภาพมีลักษณะตรงตามคลาสใดคลาสหนึ่งเท่านั้น ในมุมมองของงานนี้ MediaPipe Pose ทำหน้าที่ "
            "เป็นตัวประมาณโครงกระดูก (skeleton proxy) ส่วนอัลกอริทึมติดตามทำหน้าที่ตีความข้อมูลท่าทางให้กลายเป็นสถานะการเคลื่อนไหว การนับจำนวนครั้ง "
            "และค่า progress ซึ่งสอดคล้องกับแนวคิดของ pose-to-pose animation และ motion state transitions"
        ),
        font_size=12,
    )
    add_heading(doc, "1.2 นิยามเชิงรูปนัย (Formal Definition)", 2)
    add_bullets(
        doc,
        [
            "Input: ลำดับของ 2D joint landmarks พร้อมค่าความเชื่อมั่น/visibility จาก webcam หรือไฟล์ CSV ที่บันทึกไว้ล่วงหน้า",
            "Output: จำนวนครั้งของการเคลื่อนไหว (count), สถานะของท่า (stage), มุมที่ใช้อ้างอิงการเคลื่อนไหว (angle), และค่า progress สำหรับการแสดงผล",
            "Constraints: ต้องตอบสนองใกล้เคียงแบบ real-time, ลดผลกระทบของ occlusion และ landmark jitter, และรักษารูปแบบ output เดิมเพื่อให้ระบบ UI ใช้ต่อได้",
            "Real-time target: ให้ระบบตอบสนองต่อผู้ใช้บนหน้าเว็บแบบ interactive โดยไม่ทำให้ feedback กระตุกจนใช้งานจริงได้ยาก",
        ],
    )
    add_heading(doc, "1.3 ความท้าทาย", 2)
    add_bullets(
        doc,
        [
            "ค่ามุมของข้อต่อมี noise จากการตรวจจับ pose ทำให้ stage และ progress เปลี่ยนผิดจังหวะได้",
            "เมื่อ visibility ของแขนซ้ายและขวาใกล้เคียงกัน ระบบอาจสลับ active side ถี่เกินไปและทำให้การติดตามไม่นิ่ง",
            "หาก threshold ถูกใช้แบบตรงไปตรงมาเกินไป จะเกิด threshold chatter บริเวณจุดกลับตัวของการเคลื่อนไหว",
            "ระบบต้องออกแบบให้ reproducible และ benchmark ได้ในหลายขนาดข้อมูล ไม่ใช่เพียงแสดง demo เชิงภาพ",
        ],
    )

    add_heading(doc, "2. งานที่เกี่ยวข้อง วิธีพื้นฐาน และบริบทของ State of the Art", 1)
    add_paragraph(
        doc,
        (
            "แนวทางพื้นฐานของระบบนับท่าออกกำลังกายจาก pose landmark มักประกอบด้วยสามขั้นตอน ได้แก่ การเลือกข้างของร่างกายที่จะใช้วัด "
            "การคำนวณมุมของข้อต่อ และการกำหนด threshold เพื่อเปลี่ยน stage และนับ repetition วิธีลักษณะนี้มีข้อดีคือเรียบง่าย อธิบายผลได้ และมีต้นทุนต่อเฟรมคงที่ "
            "อย่างไรก็ตาม เมื่อใช้งานจริงกับ landmark sequence ที่มี noise และ occlusion วิธีดังกล่าวมักมีความแกว่งสูง ในขณะที่ระบบ state-of-the-art "
            "ในงานรู้จำการกระทำมนุษย์นิยมใช้ temporal neural networks, transformer models หรือ 3D pose pipelines ซึ่งให้ temporal context มากกว่าแต่มีต้นทุนสูงกว่า "
            "และไม่เหมาะกับเป้าหมายของโครงงานเชิงวิศวกรรมที่ต้องเน้นความโปร่งใสและการทำซ้ำผล"
        ),
        font_size=12,
    )
    related_table = doc.add_table(rows=1, cols=3)
    related_table.style = "Table Grid"
    hdr = related_table.rows[0].cells
    hdr[0].text = "วิธี (Method)"
    hdr[1].text = "Time Complexity"
    hdr[2].text = "ข้อจำกัด (Limitation)"
    rows = [
        ("Raw threshold rep counter", "O(1) ต่อเฟรม", "ไวต่อ noise และ threshold chatter บริเวณ turning point"),
        ("Visibility-only side selection", "O(1) ต่อเฟรม", "สลับแขนบ่อยภายใต้ partial occlusion"),
        ("Temporal neural / action model (SOTA context)", "ขึ้นกับสถาปัตยกรรมโมเดล", "ใช้ compute สูง ตีความผลยาก และเกินขอบเขตโครงงาน"),
    ]
    for a, b, c in rows:
        row = related_table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    format_table(related_table)

    add_heading(doc, "3. วิธีการที่เสนอ (Proposed Approach)", 1)
    add_heading(doc, "3.1 แนวคิดหลัก (Intuition)", 2)
    add_paragraph(
        doc,
        (
            "วิธี proposed รักษาสัญญาเชิงอินเทอร์เฟซของ baseline ไว้ทั้งหมด กล่าวคือยังให้ output แบบเดียวกัน ได้แก่ count, stage, angle และ progress "
            "แต่ปรับวิธีประมวลผลภายในให้มีหน่วยความจำระยะสั้นสำหรับติดตาม active side, smoothing state และ side-lock timer "
            "แนวคิดนี้ทำให้ระบบไม่ตัดสินใจจากเฟรมเดียวอย่างหุนหันพลันแล่น แต่ใช้ข้อมูลต่อเนื่องเชิงเวลาเพื่อลดการสลับแขนและลดความสั่นของค่ามุม"
        ),
        font_size=12,
    )
    add_heading(doc, "3.2 การออกแบบอัลกอริทึม (Pseudo-code และ Workflow)", 2)
    add_code_block(
        doc,
        [
            "Algorithm 1: Confidence-aware temporal rep tracker",
            "for each frame do",
            "    estimate visibility/confidence of left and right arms",
            "    if previously locked side is still reliable then",
            "        keep the same side to reduce arm flapping",
            "    else",
            "        select the side with the higher confidence score",
            "    end if",
            "    compute elbow angle for the selected side",
            "    update EMA-smoothed angle",
            "    update stage with hysteresis thresholds",
            "    if transition down -> up occurs then counter <- counter + 1",
            "    emit {count, stage, angle, progress}",
            "end for",
        ],
    )
    add_paragraph(
        doc,
        (
            "ในการออกแบบนี้ arm flapping หมายถึงการที่ระบบสลับไปมาระหว่างแขนซ้ายและขวาถี่เกินไปเนื่องจาก visibility ของแต่ละข้างเปลี่ยนเล็กน้อยในแต่ละเฟรม "
            "ดังนั้นหากแขนที่ lock ไว้ยังน่าเชื่อถือ ระบบจะติดตามแขนข้างเดิมต่อเพื่อลด side switching ที่ไม่จำเป็น"
        ),
        font_size=12,
    )
    add_heading(doc, "3.3 การวิเคราะห์เชิงทฤษฎี: Time, Space และ Stability", 2)
    add_bullets(
        doc,
        [
            "Time Complexity: ทั้ง baseline และ proposed มีความซับซ้อน O(1) ต่อเฟรม เนื่องจากใช้จำนวน joint คงที่",
            "Space Complexity: ทั้งสองวิธีใช้พื้นที่ O(1) เพราะเก็บ state จำนวนน้อย เช่น counter, stage, active side และค่า EMA",
            "Stability Analysis: วิธี proposed ควรลด side-switch jitter และ threshold chatter ด้วย side lock และ temporal smoothing",
            "Trade-off: การเพิ่ม confidence scoring และ EMA มี constant-factor overhead จึงอาจทำให้ช้าลงเล็กน้อยในบางสเกล แม้ให้เสถียรภาพสูงขึ้น",
        ],
    )
    add_full_width_figure(doc, WORKFLOW_IMAGE, "รูปที่ 1 ระบบงานโดยรวมของการรับภาพ ประมาณท่าทาง วิเคราะห์ motion state วัด metric และแสดงผล")

    add_heading(doc, "4. รายละเอียดการพัฒนา (Implementation Details)", 1)
    add_paragraph(
        doc,
        (
            "ระบบถูกพัฒนาด้วย Python 3.12 โดยแบ่งสถาปัตยกรรมออกเป็นสองส่วนหลัก ส่วนแรกคือ live application ประกอบด้วย Flask, OpenCV, MediaPipe และ "
            "exercise-specific modules สำหรับใช้งานบนหน้าเว็บจริง ส่วนที่สองคือ benchmarkable modules ใน `src/algorithms/` ซึ่งแยก logic ที่ใช้เปรียบเทียบ "
            "baseline/proposed ออกจากกล้องและ UI เพื่อให้สามารถทดสอบแบบ reproducible ได้อย่างชัดเจน"
        ),
        font_size=12,
    )
    impl_table = doc.add_table(rows=1, cols=2)
    impl_table.style = "Table Grid"
    impl_table.rows[0].cells[0].text = "หัวข้อ"
    impl_table.rows[0].cells[1].text = "รายละเอียด"
    impl_rows = [
        ("เทคโนโลยี", "Python 3.12, Flask, OpenCV, MediaPipe, NumPy, python-docx, python-pptx"),
        ("โครงสร้างข้อมูล", "Joint tuples, TrackingState, ExponentialAngleFilter, deterministic side-lock state"),
        ("การจัดการกรณีขอบ", "visibility drop, arm switching, threshold chatter, camera read failure"),
        ("โค้ดหลัก", "`app.py`, `camera.py`, `exercises/`, `src/algorithms/`, `experiments/benchmark.py`"),
        ("ข้อมูลจริง", "`tools/collect_mocap.py` สำหรับบันทึก landmark และ `experiments/evaluate_recorded_sequence.py` สำหรับ replay"),
    ]
    for a, b in impl_rows:
        row = impl_table.add_row().cells
        row[0].text = a
        row[1].text = b
    format_table(impl_table)

    add_heading(doc, "5. การทดลอง (Experimental Setup)", 1)
    add_heading(doc, "5.1 ชุดข้อมูล", 2)
    add_paragraph(
        doc,
        (
            "การทดลองหลักใช้ข้อมูลสังเคราะห์ (synthetic data) ที่สร้างจาก motion profile ของ elbow angle แบบคาบ พร้อมการเติม Gaussian-like noise "
            "และการสลับ visibility ซ้าย/ขวาเพื่อจำลองสภาพแวดล้อมที่ pose tracker ทำงานไม่สมบูรณ์ ข้อมูลสังเคราะห์ถูกสร้างแบบ on-demand จาก "
            "`src/algorithms/synthetic_data.py` และใช้ขนาด 1,000, 10,000 และ 100,000 เฟรมตามข้อกำหนด scaling test ของรายวิชา "
            "นอกจากนี้โครงงานยังมี pipeline สำหรับข้อมูลจริง โดยใช้ `tools/collect_mocap.py` บันทึก arm landmark sequence จาก webcam แล้วประเมินซ้ำด้วย "
            "`experiments/evaluate_recorded_sequence.py`"
        ),
        font_size=12,
    )
    add_heading(doc, "5.2 สภาพแวดล้อม", 2)
    add_bullets(
        doc,
        [
            f"ระบบปฏิบัติการ: {platform.platform()}",
            f"สถาปัตยกรรมเครื่อง: {platform.machine()}",
            f"ข้อมูล processor จาก environment: {platform.processor()}",
            "ภาษาและ runtime: Python 3.12",
            "การรัน benchmark: local execution จาก repository root",
        ],
    )
    add_heading(doc, "5.3 ตัวชี้วัด", 2)
    add_bullets(
        doc,
        [
            "Runtime (total ms) และ ms/frame สำหรับวัดประสิทธิภาพเชิงเวลา",
            "Peak traced memory (KB) สำหรับวัดการใช้หน่วยความจำเชิงสัมพัทธ์",
            "Repetition-count error เทียบกับ expected reps ของข้อมูลสังเคราะห์",
            "Side-switch jitter เพื่อวัดความแกว่งของ active side ระหว่างซ้าย/ขวา",
            "Visual quality ผ่าน side-by-side demo video และ plots ที่สร้างจาก benchmark",
        ],
    )
    exp_table = doc.add_table(rows=1, cols=3)
    exp_table.style = "Table Grid"
    exp_table.rows[0].cells[0].text = "องค์ประกอบ"
    exp_table.rows[0].cells[1].text = "ค่า"
    exp_table.rows[0].cells[2].text = "จุดประสงค์"
    exp_rows = [
        ("Synthetic sizes", "1k / 10k / 100k frames", "ทดสอบ scaling test ตาม rubric"),
        ("Expected reps", "20 / 200 / 2000", "ใช้วัด repetition-count error"),
        ("Real-data path", "record CSV -> replay evaluation", "รองรับข้อมูลจริงโดยไม่ผูกกับ webcam แบบสดตลอดเวลา"),
        ("Visualization", "plots.png + synthetic_bicep_curl.mp4", "แสดงผลเชิงภาพและ side-by-side comparison"),
    ]
    for a, b, c in exp_rows:
        row = exp_table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    format_table(exp_table)

    add_heading(doc, "6. ผลการทดลองและการวิเคราะห์ (Results & Evaluation)", 1)
    result_table = doc.add_table(rows=1, cols=8)
    result_table.style = "Table Grid"
    headers = [
        "Input Size",
        "Baseline (ms/frame)",
        "Proposed (ms/frame)",
        "Speedup",
        "Baseline Memory (KB)",
        "Proposed Memory (KB)",
        "Baseline Jitter",
        "Proposed Jitter",
    ]
    for idx, header in enumerate(headers):
        result_table.rows[0].cells[idx].text = header

    for size in sorted(grouped):
        info = speeds[size]
        baseline = info["baseline"]
        proposed = info["proposed"]
        row = result_table.add_row().cells
        row[0].text = f"{size:,}"
        row[1].text = f"{baseline.ms_per_frame:.6f}"
        row[2].text = f"{proposed.ms_per_frame:.6f}"
        row[3].text = f"{info['speedup']:.2f}%"
        row[4].text = f"{baseline.peak_kb:.3f}"
        row[5].text = f"{proposed.peak_kb:.3f}"
        row[6].text = str(baseline.side_switches)
        row[7].text = str(proposed.side_switches)
    format_table(result_table, header_size=10.5, body_size=9.5)

    add_paragraph(
        doc,
        (
            f"จากตารางพบว่า proposed method เร็วกว่า baseline ที่ขนาด 1,000 และ 10,000 เฟรม คิดเป็น {m1['speedup']:.2f}% และ {m2['speedup']:.2f}% "
            f"ตามลำดับ อย่างไรก็ตาม ที่ขนาด 100,000 เฟรม runtime เฉลี่ยต่อเฟรมเพิ่มขึ้นจาก {m3['baseline'].ms_per_frame:.6f} เป็น "
            f"{m3['proposed'].ms_per_frame:.6f} ms/frame หรือช้าลงประมาณ {abs(m3['speedup']):.2f}% สะท้อนว่า side-lock และ temporal filtering "
            "ให้เสถียรภาพที่ดีขึ้นแลกกับ constant-factor overhead เล็กน้อยที่สเกลสูงสุด"
        ),
        font_size=12,
    )
    add_paragraph(
        doc,
        (
            f"ในด้านความถูกต้อง ข้อมูลสังเคราะห์ทุกขนาดให้ repetition-count error เท่ากับศูนย์ทั้ง baseline และ proposed จึงแสดงให้เห็นว่า proposed "
            "ไม่ได้ลดทอนความถูกต้องพื้นฐานของการนับจำนวนครั้ง ขณะที่ในด้านความเสถียร proposed ลด side-switch jitter ได้ทุกสเกล "
            f"โดยที่ 100,000 เฟรมลดจาก {m3['baseline'].side_switches} เหลือ {m3['proposed'].side_switches} ครั้ง ({m3['jitter_reduction']:.2f}%) "
            "ทำให้การติดตามมีความนิ่งขึ้นและลดโอกาสที่ progress หรือ stage จะกระโดดผิดจังหวะ"
        ),
        font_size=12,
    )
    add_paragraph(
        doc,
        (
            f"ด้านหน่วยความจำ proposed ใช้ peak traced memory น้อยกว่าที่สเกล 10,000 และ 100,000 เฟรม โดยเฉพาะที่ 100,000 เฟรมลดลงจาก "
            f"{m3['baseline'].peak_kb:.3f} KB เหลือ {m3['proposed'].peak_kb:.3f} KB ({m3['memory_reduction']:.2f}%) แม้ที่สเกล 1,000 เฟรม "
            "จะมีค่า peak memory สูงขึ้นเล็กน้อยก็ตาม ผลลัพธ์ดังกล่าวบ่งชี้ว่าประโยชน์ของ proposed ชัดเจนที่สุดเมื่อ sequence ยาวขึ้นและความแกว่งเชิงเวลามีผลสะสมมากขึ้น"
        ),
        font_size=12,
    )
    add_full_width_figure(doc, PLOT_IMAGE, "รูปที่ 2 ผลเปรียบเทียบ runtime, count error, peak memory และ side-switch jitter ในทั้งสามสเกลของการทดลอง")
    add_full_width_figure(doc, DEMO_STRIP_IMAGE, "รูปที่ 3 ตัวอย่างเฟรมจาก demo แบบ side-by-side ระหว่าง baseline และ proposed สำหรับการประเมินเชิงภาพ")

    add_heading(doc, "7. อภิปรายผลโดยรวม (Discussion)", 1)
    add_paragraph(
        doc,
        (
            "ผลการทดลองสนับสนุนข้อสังเกตสำคัญสองประการ ประการแรก proposed method ช่วยลดความแกว่งของการติดตามได้จริงผ่านการลด side-switch jitter "
            "ซึ่งเป็น metric ที่เชื่อมโยงโดยตรงกับคุณภาพของ feedback บนหน้าเว็บ ประการที่สอง การเพิ่ม temporal filtering และ side locking "
            "ทำให้ระบบมี trade-off ที่สมเหตุสมผล กล่าวคือให้เสถียรภาพดีขึ้นอย่างสม่ำเสมอ ขณะที่ต้นทุนด้านเวลาเพิ่มขึ้นเพียงเล็กน้อยในกรณีข้อมูลขนาดใหญ่มาก"
        ),
        font_size=12,
    )
    add_paragraph(
        doc,
        (
            "เมื่อเทียบกับบริบทของ state-of-the-art วิธีที่เสนอไม่ได้แข่งขันกับโมเดลเชิงลึกขนาดใหญ่ในเชิงความซับซ้อนของ representation "
            "แต่มีจุดแข็งด้าน explainability, reproducibility, และความเหมาะสมต่องาน interactive engineering artifact ภายในวิชา Computer Animation "
            "ซึ่งให้ความสำคัญกับความสัมพันธ์ระหว่าง pose, state transition และ visual feedback มากกว่าการใช้โมเดล black-box"
        ),
        font_size=12,
    )
    add_bullets(
        doc,
        [
            "จุดแข็ง: สถาปัตยกรรมเรียบง่าย อธิบายผลได้ และมี baseline/proposed comparison ที่ชัดเจน",
            "จุดแข็ง: มี scaling test, plots, demo video, tests และ pipeline สำหรับข้อมูลจริง",
            "ข้อจำกัด: benchmark เชิงปริมาณครอบคลุม motion primitive หลักเพียงชนิดเดียว (bicep curl) ไม่ใช่ทุก exercise mode",
            "ข้อจำกัด: ไม่มี human mocap sequence ที่ commit ลง repo โดยตรง เนื่องจากประเด็นด้าน privacy และการบันทึกขึ้นกับผู้ใช้",
        ],
    )

    add_heading(doc, "8. บทสรุป (Conclusion)", 1)
    add_paragraph(
        doc,
        (
            "โครงงานนี้นำเสนอระบบติดตามการเคลื่อนไหวของท่าออกกำลังกายแบบเรียลไทม์ที่ตีความปัญหาในกรอบของ Computer Animation อย่างชัดเจน ผ่านการใช้ "
            "pose skeleton, key poses, interpolation ของ progress และ temporal state transitions วิธี proposed ยังคงรูปแบบ output เดิมของ baseline "
            "ไว้ทั้งหมด แต่ปรับกลไกภายในเพื่อเพิ่มเสถียรภาพของการติดตาม ผลลัพธ์เชิงทดลองแสดงให้เห็นว่าแนวทางดังกล่าวช่วยลด side-switch jitter และลดหน่วยความจำที่สเกลสูง "
            "พร้อมทั้งรักษาความถูกต้องของการนับท่าได้อย่างครบถ้วน แม้จะมี runtime overhead เล็กน้อยในกรณีข้อมูลขนาดใหญ่ที่สุด จึงนับเป็น engineering artifact "
            "ที่มีคุณค่าทั้งในเชิงรายวิชาและการนำไปใช้ต่อใน portfolio"
        ),
        font_size=12,
    )

    add_heading(doc, "9. Demo, Reproducibility และโครงสร้างไฟล์ที่ส่ง", 1)
    add_paragraph(
        doc,
        (
            "เพื่อให้เป็นไปตาม rubric ด้าน reproducibility และ demo โครงงานนี้จัดเตรียมทั้ง live web application, benchmark scripts, test suite, generated plots, "
            "และ side-by-side demo video ผู้ตรวจสามารถรันระบบหลักด้วย `python src/main.py --mode web` รัน benchmark ด้วย "
            "`python src/main.py --mode benchmark` และรัน demo video ด้วย `python src/main.py --mode demo` นอกจากนี้ยังสามารถอัดข้อมูลจริงผ่าน "
            "`python tools/collect_mocap.py --output data/mocap/recorded_arm_sequence.csv` แล้ว replay เพื่อประเมินผลซ้ำด้วย "
            "`python experiments/evaluate_recorded_sequence.py data/mocap/recorded_arm_sequence.csv --expected-reps 10`"
        ),
        font_size=12,
    )
    reproducibility_table = doc.add_table(rows=1, cols=2)
    reproducibility_table.style = "Table Grid"
    reproducibility_table.rows[0].cells[0].text = "องค์ประกอบ"
    reproducibility_table.rows[0].cells[1].text = "ไฟล์ / คำสั่ง"
    repro_rows = [
        ("รายงาน", "report/final_report_formal.docx"),
        ("สไลด์", "report/final_presentation.pptx / report/final_presentation_th.pptx"),
        ("ผล benchmark", "experiments/results/runtime.csv และ experiments/results/plots.png"),
        ("demo video", "demo/synthetic_bicep_curl.mp4"),
        ("unit tests", "python -m unittest tests.test_simulation"),
    ]
    for a, b in repro_rows:
        row = reproducibility_table.add_row().cells
        row[0].text = a
        row[1].text = b
    format_table(reproducibility_table)

    add_heading(doc, "เอกสารอ้างอิง (References)", 1)
    references = [
        "[1] G. Bradski, “The OpenCV Library,” Dr. Dobb’s Journal of Software Tools, 2000.",
        "[2] V. Bazarevsky et al., “BlazePose: On-device Real-time Body Pose Tracking,” arXiv:2006.10204, 2020.",
        "[3] Google MediaPipe Team, “MediaPipe Pose Documentation,” official developer documentation.",
        "[4] G. Casiez, N. Roussel, and D. Vogel, “1 Euro Filter: A Simple Speed-based Low-pass Filter for Noisy Input in Interactive Systems,” CHI 2012.",
        "[5] R. Parent, Computer Animation: Algorithms and Techniques, 3rd ed., Morgan Kaufmann, 2012.",
        "[6] A. Menache, Understanding Motion Capture for Computer Animation and Video Games, 2nd ed., Morgan Kaufmann, 2010.",
    ]
    for ref in references:
        add_paragraph(doc, ref, font_size=11)

    add_heading(doc, "ภาคผนวก A. โครงสร้างโฟลเดอร์ที่ใช้ส่งงาน", 1)
    add_code_block(
        doc,
        [
            "computer-animation-project-/",
            "├── README.md",
            "├── report/",
            "│   ├── final_report_formal.docx",
            "│   ├── final_presentation.pptx",
            "│   └── final_presentation_th.pptx",
            "├── src/",
            "├── experiments/",
            "├── data/",
            "├── visualization/",
            "├── demo/",
            "├── tests/",
            "├── tools/",
            "├── app.py",
            "├── camera.py",
            "└── requirements.txt",
        ],
    )

    add_heading(doc, "ภาคผนวก B. หมายเหตุเกี่ยวกับข้อมูลสังเคราะห์", 1)
    add_paragraph(
        doc,
        (
            "ข้อมูลสังเคราะห์ในงานนี้ไม่ได้ถูกเก็บเป็นไฟล์ขนาดใหญ่ใน repository แต่ถูกสร้างแบบ on-demand จาก `src/algorithms/synthetic_data.py` "
            "เพื่อให้ benchmark ทำซ้ำได้ทุกครั้งโดยใช้ parameter เดิม ทั้งนี้ synthetic sequence ประกอบด้วย elbow-angle motion profile, noise injection, "
            "visibility imbalance, และ dropout segments เพื่อเลียนแบบความไม่สมบูรณ์ของ pose tracking ในการใช้งานจริง"
        ),
        font_size=12,
    )


def build_report():
    doc = Document()
    set_document_language(doc)

    cover = doc.sections[0]
    set_section_margins(cover, margin_inches=0.8)
    set_section_columns(cover, columns=1)

    _, speeds = compute_metrics()
    add_cover_page(doc, speeds)

    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_margins(body, margin_inches=0.6)
    set_section_columns(body, columns=2)

    add_report_body(doc)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Generated {OUTPUT_FILE}")


if __name__ == "__main__":
    build_report()
